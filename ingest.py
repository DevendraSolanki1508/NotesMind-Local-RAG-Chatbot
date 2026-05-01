import os
import sys
import fitz  # PyMuPDF
import chromadb
import pytesseract
import pdf2image
from pathlib import Path
from llama_index.core import Document, VectorStoreIndex, Settings, StorageContext
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore

# ── Paths ──────────────────────────────────────────────────────
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\poppler-25.12.0\Library\bin"

DOCS_DIR   = "./docs"
CHROMA_DIR = "./chroma_db"
COLLECTION = "notesmind"

# ── Embedding model ────────────────────────────────────────────
print("Loading embedding model...")
embed_model = HuggingFaceEmbedding(
    model_name="BAAI/bge-small-en-v1.5",
    embed_batch_size=32
)
Settings.embed_model = embed_model
Settings.llm = None
Settings.node_parser = SentenceSplitter(chunk_size=256, chunk_overlap=32)

# ── OCR function ───────────────────────────────────────────────
def extract_text_with_ocr(pdf_path: str) -> str:
    print(f"  Running OCR on: {Path(pdf_path).name}")
    pages = pdf2image.convert_from_path(
        pdf_path,
        dpi=200,
        poppler_path=POPPLER_PATH
    )
    text = ""
    for i, page in enumerate(pages):
        print(f"    Page {i+1}/{len(pages)}...")
        text += pytesseract.image_to_string(page, lang="eng")
    return text

# ── Document loader ────────────────────────────────────────────
def load_documents(docs_dir: str):
    documents = []
    files = os.listdir(docs_dir)

    if not files:
        print("ERROR: No files found in ./docs")
        sys.exit(1)

    for filename in files:
        filepath = os.path.join(docs_dir, filename)
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            # Try text extraction first
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()

            # If barely any text, it's a scanned PDF
            if len(text.strip()) < 100:
                print(f"Scanned PDF detected: {filename} — running OCR...")
                text = extract_text_with_ocr(filepath)
            else:
                print(f"Text PDF loaded: {filename}")

        elif ext in [".txt", ".md"]:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            print(f"Loaded: {filename}")

        elif ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            print(f"Loaded: {filename}")

        else:
            print(f"Skipping unsupported file: {filename}")
            continue

        if text.strip():
            documents.append(Document(
                text=text,
                metadata={"filename": filename}
            ))

    return documents

# ── Main ───────────────────────────────────────────────────────
print("Reading documents...")
documents = load_documents(DOCS_DIR)
print(f"\nSuccessfully loaded {len(documents)} document(s).")

# ── ChromaDB setup ─────────────────────────────────────────────
chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)
try:
    chroma_client.delete_collection(COLLECTION)
    print("Cleared old index.")
except Exception:
    pass

collection   = chroma_client.get_or_create_collection(COLLECTION)
vector_store = ChromaVectorStore(chroma_collection=collection)
storage_ctx  = StorageContext.from_defaults(vector_store=vector_store)

# ── Build index ────────────────────────────────────────────────
print("\nBuilding index...")
index = VectorStoreIndex.from_documents(
    documents,
    storage_context=storage_ctx,
    show_progress=True
)

print("\n✓ Done! Indexed successfully.")
print("Next step — run:  streamlit run app.py")